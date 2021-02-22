#!/usr/bin/env python3

import docx
from docx.shared import Cm
import docx2txt
import pandas as pd
import os
from gauge import gauge
import copy

def image_lookup(doc):
    # Adapted from https://stackoverflow.com/a/61331396

    lookup = {}

    img_path = doc.replace('.docx','')
    os.makedirs(img_path, exist_ok=True)

    # Extract the images to img_folder/
    docx2txt.process(doc, img_path)

    # Open your .docx document
    doc = docx.Document(doc)

    # Save all 'rId:filenames' relationships in an dictionary named rels
    rels = {}
    for r in doc.part.rels.values():
        if isinstance(r._target, docx.parts.image.ImagePart):
            rels[r.rId] = os.path.basename(r._target.partname)

    # Then process your text
    for paragraph in doc.paragraphs:
        # If you find an image
        if paragraph.text.strip():
            bits = paragraph.text.strip().split("\t")
            number = int(bits[0])
        if 'Graphic' in paragraph._p.xml:
            # Get the rId of the image
            for rId in rels:
                if f'"{rId}"' in paragraph._p.xml:
                    # Your image will be in os.path.join(img_path, rels[rId])
                    #print(rels[rId])
                    lookup[number] = os.path.join(img_path, rels[rId])
    return lookup

os.makedirs("output", exist_ok=True)

template = docx.Document('input/Personalised report template sent to Nick e centre.docx')
doc2_map = image_lookup("input/Required document 2 - Knowledge about each technology.docx")
doc3_map = image_lookup("input/Required document 3 - Technology implementation level.docx")
doc4_map = image_lookup("input/Required document 4 - Technology readiness level on each indicator.docx")

doc5 = docx.Document("input/Required document 5 - Weakness and improvements on your technology readiness.docx")
doc5.paragraphs[0].runs[0].text = "1\tAhmed" # Fix weirdness at start of document
doc5_lookup = {}
for para in doc5.paragraphs:
    if para.runs and para.runs[0].font.highlight_color:
        bits = para.text.split("\t")
        number = int(bits[0])
    else:
        if number not in doc5_lookup:
            doc5_lookup[number] = []
        doc5_lookup[number].append(para)

df = pd.read_excel("input/Required document 1 - Data spreadsheet for e-research center.xlsx")
print(df)

LEVELS = ['Low','Low to medium','Medium','Medium to High','High']

for i, row in df.iterrows():
    doc = copy.deepcopy(template)
    if i == 0:
        continue
    number = row[0]
    V1 = row["V1"]
    V2 = str(round(row["V2"], 1))
    V3 = row["V3"]
    V4 = str(round(row["V4 (constant)"], 1))
    print(number, V1, V2, V3, V4)
    filename = f'output/{number}_{V1}'
    image_filename = filename + ".png"
    gauge(labels=['Low', '', 'Medium', '', 'High'],
          colors='RdYlGn',
          arrow=row["V2"] / 2,
          fname=image_filename)
    for para in doc.paragraphs:
        if "…V1" in para.text:
            para.text = para.text.replace("…V1", V1)
        if "***V2" in para.text:
            para.text = para.text.replace("***V2", V2)
        if "***V3" in para.text:
            para.text = para.text.replace("***V3", V3)
        if para.text.strip() == "Diagram to show the score of the company on average (V2) and the score of the industry on average (V4).":
            para.text = ""
            run = para.add_run()
            run.add_picture(image_filename)
        if para.text.strip() == "Diagram from document 2":
            para.text = ""
            run = para.add_run()
            run.add_picture(doc2_map[number])
        if para.text.strip() == "Diagram from document 3":
            para.text = ""
            run = para.add_run()
            run.add_picture(doc3_map[number])
        if para.text.strip() == "Diagram from document 4":
            para.text = ""
            run = para.add_run()
            run.add_picture(doc4_map[number], width=Cm(16))
        if para.text.strip() == "Texts from document 5":
            para.text = ""
            doc5_paras = doc5_lookup[number]
            for other_para in doc5_paras:
                para.insert_paragraph_before(other_para.text, style=other_para.style)
    doc.save(filename + ".docx")
    print(f"{filename} saved")